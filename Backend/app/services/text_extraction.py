import os
from pypdf import PdfReader
from docx import Document as DocxDocument

class TextExtractionServices:

    SUPPORTED_FORMAT = {
            'text/plain': ['txt', 'md'],
            'application/pdf': ['pdf'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['docx']
    }



    def _extract_text_file(self, file_path: str)->str:
        """extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text.strip()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            return text.strip()
        

    def is_supported(self, content_type: str, file_path: str ) -> bool:

        extension = self._get_extension(file_path)
        
        if content_type in self.SUPPORTED_FORMAT:
            return True

        for mime_type, extensions in self.SUPPORTED_FORMAT.items():
            if extension in extensions:
                return True

        return False
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF files using pypdf"""

        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            full_text = '\n\n'.join(text_parts)
            return full_text.strip()
        
        except Exception as e:
            raise Exception(f"Failed to extract text from pdf: {str(e)}")
        

    def _extract_docx(self, file_path: str)-> str:
        """Extract text from word document"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                    
                full_text = '\n\n'.join(text_parts)
                return full_text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract from the docx: {str(e)}")


    def _get_extension(self, file_path: str)-> str:
        """getting the file extension without dot"""
        return os.path.splitext(file_path)[1].lstrip('.').lower()
    
    def extract_text(self, file_path: str, content_type: str) -> str:
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found:{file_path}")
        
        extension = self._get_extension(file_path)

        if content_type == 'text/plain' or extension in ['txt', 'md']:
            return self._extract_text_file(file_path)
        
        elif content_type == 'application/pdf' or extension == 'pdf':
            return self._extract_pdf(file_path)
        
        elif 'wordpreocessingml' in content_type or extension =='docx':
            return self._extract_docx(file_path)
        
        else:
            raise ValueError(

                f"Unsupported file format: {content_type} (extension: {extension})."
                f"Supporteed format : .txt, .md, .pdf, .docx"
            )
        
# a = TextExtractionServices()
# a.extract_text("/Users/divyansh/Downloads/Hammer.pdf", 'string')